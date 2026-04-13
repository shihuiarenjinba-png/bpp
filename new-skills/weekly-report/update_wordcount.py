import sys
import os
from docx import Document

def count_and_update_wordcount(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found.")
        return False

    try:
        doc = Document(file_path)
        
        # 1. 全テキストを抽出して文字数をカウント（スペースや改行を除く）
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text
        
        # 文字数カウント（空白文字を除外する場合の例。必要に応じて調整可能）
        # ここでは単純に全文字数をカウントします
        char_count = len(full_text.replace(" ", "").replace("\n", "").replace("\r", ""))
        
        # 2. プレースホルダ「XXX文字」を探して置換する
        # 段落内のテキストをスキャン
        found = False
        for para in doc.paragraphs:
            if "XXX文字" in para.text:
                para.text = para.text.replace("XXX文字", f"{char_count}文字")
                found = True
        
        # テーブル内のテキストもスキャン（念のため）
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if "XXX文字" in para.text:
                                para.text = para.text.replace("XXX文字", f"{char_count}文字")
                                found = True

        if found:
            doc.save(file_path)
            print(f"Successfully updated word count: {char_count} characters.")
            return True
        else:
            print("Placeholder 'XXX文字' not found in the document.")
            return False

    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 update_wordcount.py <path_to_docx>")
    else:
        count_and_update_wordcount(sys.argv[1])
