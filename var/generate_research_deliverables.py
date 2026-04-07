from pathlib import Path
import subprocess
from textwrap import dedent

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

import generate_thesis_progress_ppt as ppt_builder
from thesis_research_state import (
    BACKGROUND,
    CHAPTER_STATUS,
    CODE_OUTPUT_PATH,
    CONTRIBUTIONS,
    CURRENT_ISSUES,
    CURRENT_PROGRESS,
    DATA_PLAN,
    DOCX_OUTPUT_PATH,
    HYPOTHESES,
    MASTER_PLAN_PATH,
    METHODS,
    PPT_OUTPUT_PATH,
    PURPOSES,
    REFERENCE_FILES,
    RESEARCH_DIR,
    ROADMAP,
    SHORT_TITLE,
    SUMMARY,
    THESIS_TITLE,
    THREE_LAYER_STRUCTURE,
    VARIABLE_CATALOG,
    WEB_SOURCES,
    XLSX_OUTPUT_PATH,
)


HEADER_FILL = PatternFill("solid", fgColor="222222")
SUB_FILL = PatternFill("solid", fgColor="E9E9E9")
NOTE_FILL = PatternFill("solid", fgColor="F7F7F7")
HEADER_FONT = Font(name="IPAexGothic", bold=True, color="FFFFFF")
TITLE_FONT = Font(name="IPAexGothic", bold=True, size=13, color="1F1F1F")
BODY_FONT = Font(name="IPAexGothic", size=10.5, color="1F1F1F")
LINK_FONT = Font(name="Aptos", size=10.5, color="0563C1", underline="single")
MUTED_FONT = Font(name="IPAexGothic", size=10, italic=True, color="666666")
THIN = Side(style="thin", color="D0D0D0")
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_cell(cell, fill=None, font=None, align=None):
    cell.border = BOX
    cell.font = font or BODY_FONT
    cell.alignment = align or Alignment(vertical="top", wrap_text=True)
    if fill:
        cell.fill = fill


def add_header(ws, row, values):
    for idx, value in enumerate(values, start=1):
        cell = ws.cell(row=row, column=idx, value=value)
        style_cell(
            cell,
            fill=HEADER_FILL,
            font=HEADER_FONT,
            align=Alignment(horizontal="center", vertical="center", wrap_text=True),
        )


def add_rows(ws, start_row, rows):
    for r, row in enumerate(rows, start=start_row):
        for c, value in enumerate(row, start=1):
            cell = ws.cell(row=r, column=c, value=value)
            style_cell(cell)


def set_col_widths(ws, widths):
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def setup_sheet(ws, freeze="A2", filter_ref=None):
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = True
    if filter_ref:
        ws.auto_filter.ref = filter_ref


def add_link(ws, row, col, text, url):
    cell = ws.cell(row=row, column=col, value=text)
    cell.hyperlink = url
    style_cell(cell, font=LINK_FONT)


def build_readme(ws):
    ws.title = "00_README"
    ws["A1"] = SHORT_TITLE
    ws["B1"] = "このExcelは発表資料ではなく、研究用データベースとして使う前提です。"
    style_cell(ws["A1"], fill=SUB_FILL, font=TITLE_FONT)
    style_cell(ws["B1"], fill=SUB_FILL, font=MUTED_FONT)

    rows = [
        ("主データ", MASTER_PLAN_PATH.name),
        ("使い分け", "PowerPoint は短時間説明用。Excel は情報蓄積・比較・更新用。"),
        ("このブックでやること", "計画書の要点整理、文献・ネット情報の保存、変数辞書、データ取得方針、検定設計、重点論点、進捗管理。"),
        ("このブックでやらないこと", "スライドそのものの文章をそのまま重複させること。"),
        ("最優先タスク", "長期の正規データ取得と再走。"),
    ]
    for i, (left, right) in enumerate(rows, start=3):
        style_cell(ws.cell(i, 1, left), fill=NOTE_FILL, font=TITLE_FONT)
        style_cell(ws.cell(i, 2, right))

    ws["A10"] = "シートの見方"
    style_cell(ws["A10"], fill=SUB_FILL, font=TITLE_FONT)
    guide = [
        "01_PLAN_MASTER: 最終計画書の中核要点を抜き出した整理用シート",
        "02_HYPOTHESES_TESTS: 仮説と、どの検定で何を示すかの対応表",
        "03_LOCAL_FILES_DB: ローカル資料の台帳",
        "04_WEB_SOURCES_DB: ネットで確認した論文・データソースの台帳",
        "05_VARIABLE_DICTIONARY: 変数定義、系列ID、役割、変換方針",
        "06_DATA_ACQUISITION_PLAN: 実データ取得と再走のための具体手順",
        "07_LITERATURE_INSIGHTS: 読み取れた学術的示唆の蓄積",
        "08_CPI_RMW_FOCUS: 重点仮説の専用メモ",
        "09_PROGRESS_TASKBOARD: 現在の進捗・課題・次の行動",
        "10_FREE_NOTES: 自由記述",
    ]
    for idx, text in enumerate(guide, start=11):
        style_cell(ws.cell(idx, 2, text))
        style_cell(ws.cell(idx, 1, f"{idx-10:02d}"), fill=NOTE_FILL, font=TITLE_FONT, align=Alignment(horizontal="center", vertical="center"))
    set_col_widths(ws, {"A": 18, "B": 96})
    setup_sheet(ws, freeze="A2")


def build_plan_master(ws):
    ws.title = "01_PLAN_MASTER"
    add_header(ws, 1, ["区分", "内容", "研究上の意味"])
    rows = []
    rows.extend([["背景", item, "問題設定の起点"] for item in BACKGROUND])
    rows.extend([["研究目的", item, "最終計画書の目的節に対応"] for item in PURPOSES])
    rows.extend([["データ方針", item, "実証前提"] for item in DATA_PLAN])
    rows.extend([["期待される貢献", item, "結論の軸"] for item in CONTRIBUTIONS])
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 16, "B": 70, "C": 24})
    setup_sheet(ws, filter_ref=f"A1:C{len(rows)+1}")


def build_hypotheses(ws):
    ws.title = "02_HYPOTHESES_TESTS"
    add_header(ws, 1, ["仮説ID", "仮説名", "内容", "使う検定", "主な図表", "判定のポイント"])
    test_map = {
        "H1": ("グレンジャー因果性、ブロック外生性、IRF", "因果性一覧表、IRF図", "金融市場→マクロ、マクロ→金融の両方向が見えるか"),
        "H2": ("グレンジャー因果性、OOS R²、方向的中率", "遅行ファクター候補表、比較表", "特定ファクターの遅行特性と予測力が対応するか"),
        "H3": ("CPI→RMW 個別ラグ検定、多重検定補正", "ラグ別有意性表、方向的中率表", "補正後も頑健に残るか"),
    }
    rows = []
    for h in HYPOTHESES:
        tests, figs, point = test_map[h["id"]]
        rows.append([h["id"], h["title"], h["detail"], tests, figs, point])
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 10, "B": 24, "C": 48, "D": 28, "E": 22, "F": 30})
    setup_sheet(ws, filter_ref="A1:F4")


def build_local_files(ws):
    ws.title = "03_LOCAL_FILES_DB"
    add_header(ws, 1, ["ID", "資料名", "区分", "優先度", "役割", "このExcelでの使い道"])
    rows = [[item["id"], item["name"], "Local file", item["priority"], item["role"], item["note"]] for item in REFERENCE_FILES]
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 8, "B": 56, "C": 12, "D": 10, "E": 20, "F": 44})
    setup_sheet(ws, filter_ref=f"A1:F{len(rows)+1}")


def build_web_sources(ws):
    ws.title = "04_WEB_SOURCES_DB"
    add_header(ws, 1, ["区分", "名前", "URL", "何に使うか", "備考"])
    for idx, item in enumerate(WEB_SOURCES, start=2):
        style_cell(ws.cell(idx, 1, item["category"]))
        style_cell(ws.cell(idx, 2, item["name"]))
        add_link(ws, idx, 3, item["url"], item["url"])
        style_cell(ws.cell(idx, 4, "公式データソースまたは学術論文の確認用"))
        style_cell(ws.cell(idx, 5, item["note"]))
    set_col_widths(ws, {"A": 16, "B": 40, "C": 54, "D": 28, "E": 40})
    setup_sheet(ws, filter_ref=f"A1:E{len(WEB_SOURCES)+1}")


def build_variable_dict(ws):
    ws.title = "05_VARIABLE_DICTIONARY"
    add_header(ws, 1, ["系列名", "区分", "配置層", "主ソース", "頻度", "変換案", "役割メモ"])
    add_rows(ws, 2, VARIABLE_CATALOG)
    set_col_widths(ws, {"A": 16, "B": 12, "C": 18, "D": 18, "E": 12, "F": 20, "G": 38})
    setup_sheet(ws, filter_ref=f"A1:G{len(VARIABLE_CATALOG)+1}")


def build_data_plan(ws):
    ws.title = "06_DATA_ACQUISITION_PLAN"
    add_header(ws, 1, ["順番", "作業", "対象", "出力", "チェックポイント", "状態"])
    rows = [
        [1, "FF 5ファクターの長期月次データを取得", "Kenneth French Data Library", "長期の factor csv/xlsx", "Mkt-RF〜CMA が揃うか", "未着手"],
        [2, "IP・CPI・VIX・為替・原油・スプレッドを取得", "FRED 各系列", "月次 macro データ", "開始時点が十分長いか", "未着手"],
        [3, "日次系列を月次にそろえる", "VIX, 為替, 原油など", "月次統一データ", "集計ルールを明記したか", "未着手"],
        [4, "非定常系列の差分化・標準化", "スプレッド, VIX 等", "分析用入力データ", "見せかけの相関を減らせるか", "未着手"],
        [5, "2006年以降239か月版と長期版を比較", "短期/長期サンプル", "比較メモ", "結果の安定性が見えるか", "未着手"],
        [6, "再走ログを残す", "全工程", "再現メモ", "次回も同じ処理ができるか", "未着手"],
    ]
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 8, "B": 26, "C": 28, "D": 22, "E": 28, "F": 10})
    setup_sheet(ws, filter_ref="A1:F7")


def build_literature_insights(ws):
    ws.title = "07_LITERATURE_INSIGHTS"
    add_header(ws, 1, ["出典", "拾えた知見", "本研究への使い方", "注意点"])
    rows = [
        ["最終計画書", "研究の中心は『ML予測力の源泉解明』であり、単純な予測競争ではない。", "PowerPoint の中心メッセージに使う。", "他資料より優先して反映する。"],
        ["最終計画書", "当初の4層構造ではなく、先行層・同時双方向層・遅行層の3層構造で再構築する。", "モデル構造の説明に使う。", "古い4層説明を混ぜない。"],
        ["最終計画書", "CPI→RMW を重点仮説として扱う。", "重点分析と図表配置に使う。", "複数ラグと補正後有意性が鍵。"],
        ["Gu, Kelly, Xiu (2020)", "NN3 が月次 stock-level OOS R2 0.40% と最良。木系と NN が強い。", "ML側の問題提起と背景に使う。", "0.40% は R2 であり『40%のリスク』ではない。"],
        ["Gu, Kelly, Xiu (2020)", "有力シグナルは momentum・liquidity・volatility・valuation。", "補助的な候補ファクター整理に使う。", "本研究の本筋はVARでの因果性分析。"],
        ["Welch & Goyal (2008)", "多くの予測変数はヒストリカル平均に負ける。", "OOS 評価を過大に期待しない前提整理に使う。", "部分的な予測力でも意味があると整理する。"],
        ["De Oliveira et al. (2020)", "CPI イノベーションが収益性プレミアムに関係する可能性。", "CPI→RMW 仮説の補強に使う。", "時系列の因果性としては本研究で掘り下げる。"],
    ]
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 24, "B": 44, "C": 34, "D": 30})
    setup_sheet(ws, filter_ref="A1:D8")


def build_cpi_rmw(ws):
    ws.title = "08_CPI_RMW_FOCUS"
    add_header(ws, 1, ["項目", "内容", "今後確認すること"])
    rows = [
        ["位置づけ", "最終計画書で最も具体的に書かれている重点仮説。", "他の発見より先に専用表を作る。"],
        ["仮説内容", "CPI は多重検定補正後も RMW に対して頑健な先行性を持つ。", "1, 3, 6, 12か月ラグを比較する。"],
        ["理論的背景", "収益性プレミアムはインフレ期待や価格転嫁の差を通じて影響を受ける可能性がある。", "De Oliveira et al. (2020) の関連記述を補強する。"],
        ["必要な検定", "個別ラグのグレンジャー因果性、多重検定補正、方向的中率、IRF。", "重点図表を先に設計する。"],
        ["期待する出力", "ラグ別有意性表、補正後有意性、IRF、方向的中率。", "論文の核となる図表セットにする。"],
        ["失敗した場合の整理", "有意でない場合も、サンプル期間・変換・逆方向因果性を整理して限界として書く。", "否定結果でも学術的な意味を残す。"],
    ]
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 16, "B": 54, "C": 34})
    setup_sheet(ws, filter_ref="A1:C7")


def build_progress(ws):
    ws.title = "09_PROGRESS_TASKBOARD"
    add_header(ws, 1, ["区分", "状態", "項目", "内容", "次の行動"])
    rows = []
    for item in CURRENT_PROGRESS:
        rows.append(["進捗", "確認済み", "現在の到達点", item, "維持"])
    for item in CURRENT_ISSUES:
        rows.append(["課題", "未解決", "ボトルネック", item, "解消策を具体化"])
    for chapter, title, status, note in CHAPTER_STATUS:
        rows.append(["章進捗", status, chapter, f"{title}: {note}", "本文更新"])
    rows.extend([
        ["次工程", "未着手", "長期データ取得", "2006年以降より長いサンプルを確保して再走する。", "データソースから取得開始"],
        ["次工程", "未着手", "3層VAR再構築", "4層案ではなく3層案で推定式を整理する。", "入力系列を確定"],
        ["次工程", "未着手", "重点検定", "CPI→RMW のラグ別検定と補正後評価を行う。", "検定コード準備"],
    ])
    add_rows(ws, 2, rows)
    set_col_widths(ws, {"A": 10, "B": 12, "C": 18, "D": 56, "E": 22})
    setup_sheet(ws, filter_ref=f"A1:E{len(rows)+1}")


def build_notes(ws):
    ws.title = "10_FREE_NOTES"
    add_header(ws, 1, ["日付", "分類", "メモ", "出典", "反映先", "優先度", "未処理論点", "備考"])
    for row in range(2, 26):
        for col in range(1, 9):
            style_cell(ws.cell(row, col, ""))
    ws["A27"] = "メモ"
    ws["B27"] = "新しい論文・コメント・先生からの修正点はまずここに置き、必要に応じて他シートに転記する。"
    style_cell(ws["A27"], fill=SUB_FILL, font=TITLE_FONT)
    style_cell(ws["B27"], fill=SUB_FILL, font=MUTED_FONT)
    set_col_widths(ws, {"A": 12, "B": 14, "C": 42, "D": 20, "E": 16, "F": 10, "G": 20, "H": 18})
    setup_sheet(ws, filter_ref="A1:H25")


def build_excel():
    wb = Workbook()
    wb.template = False
    build_readme(wb.active)
    build_plan_master(wb.create_sheet())
    build_hypotheses(wb.create_sheet())
    build_local_files(wb.create_sheet())
    build_web_sources(wb.create_sheet())
    build_variable_dict(wb.create_sheet())
    build_data_plan(wb.create_sheet())
    build_literature_insights(wb.create_sheet())
    build_cpi_rmw(wb.create_sheet())
    build_progress(wb.create_sheet())
    build_notes(wb.create_sheet())
    wb.save(XLSX_OUTPUT_PATH)


def set_doc_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "IPAexGothic"
    normal.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "IPAexGothic"
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 3"].font.size = Pt(11.5)


def add_doc_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table


def build_consolidated_docx():
    doc = Document()
    set_doc_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(THESIS_TITLE)
    run.bold = True
    run.font.name = "IPAexGothic"
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("統合進捗レポート")
    run.font.name = "IPAexGothic"
    run.font.size = Pt(12)

    doc.add_paragraph(f"主データ: {MASTER_PLAN_PATH.name}")
    doc.add_paragraph("この文書は、最終計画書を主軸に、補助資料と確認済みの外部情報を統合した整理版です。")

    doc.add_heading("1. 文書の位置づけ", level=1)
    add_doc_bullets(doc, [
        "最終計画書を最優先の基準文書とし、他のローカル資料は補足情報として扱う。",
        "PowerPoint は短時間説明用、Excel は研究データベース用、本 Word 文書は統合読み物として使う。",
        "研究の主軸は、MLの高い個別株予測力の源泉を、FFファクターとマクロ経済変数の因果性から説明することにある。",
    ])

    doc.add_heading("2. 研究背景", level=1)
    add_doc_bullets(doc, BACKGROUND)

    doc.add_heading("3. 研究目的", level=1)
    add_doc_bullets(doc, PURPOSES)

    doc.add_heading("4. 研究仮説", level=1)
    add_doc_table(
        doc,
        ["仮説ID", "仮説名", "内容"],
        [[h["id"], h["title"], h["detail"]] for h in HYPOTHESES],
    )

    doc.add_heading("5. モデル構造", level=1)
    doc.add_paragraph("最終計画書では、当初の4層案を修正し、以下の3層構造で再構築する方針が示されています。")
    add_doc_table(
        doc,
        ["層", "構成変数", "役割", "経済的解釈"],
        [[x["layer"], x["variables"], x["role"], x["interpretation"]] for x in THREE_LAYER_STRUCTURE],
    )

    doc.add_heading("6. データと変数方針", level=1)
    add_doc_bullets(doc, DATA_PLAN)
    doc.add_heading("6.1 主要変数一覧", level=2)
    add_doc_table(
        doc,
        ["系列名", "区分", "配置層", "主ソース", "頻度", "変換案", "役割メモ"],
        VARIABLE_CATALOG,
    )

    doc.add_heading("7. 分析手法", level=1)
    add_doc_bullets(doc, METHODS)

    doc.add_heading("8. 期待される貢献", level=1)
    add_doc_bullets(doc, CONTRIBUTIONS)

    doc.add_heading("9. 現在の進捗", level=1)
    add_doc_bullets(doc, CURRENT_PROGRESS)
    doc.add_heading("9.1 現在の課題", level=2)
    add_doc_bullets(doc, CURRENT_ISSUES)

    doc.add_heading("10. 章進捗", level=1)
    add_doc_table(
        doc,
        ["章", "章題", "現在地", "現状メモ"],
        CHAPTER_STATUS,
    )

    doc.add_heading("11. 今後のロードマップ", level=1)
    for idx, item in enumerate(ROADMAP, start=1):
        doc.add_paragraph(f"{idx}. {item}")

    doc.add_heading("12. 参照資料一覧", level=1)
    add_doc_table(
        doc,
        ["ID", "資料名", "役割", "優先度", "メモ"],
        [[item["id"], item["name"], item["role"], item["priority"], item["note"]] for item in REFERENCE_FILES],
    )

    doc.add_heading("13. 外部確認情報", level=1)
    doc.add_paragraph("以下は、研究設計・データ取得・先行研究整理のために確認した主要な外部情報です。")
    rows = []
    for item in WEB_SOURCES:
        rows.append([item["category"], item["name"], item["url"], item["note"]])
    add_doc_table(doc, ["区分", "名前", "URL", "備考"], rows)

    doc.add_heading("14. 補足メモ", level=1)
    add_doc_bullets(doc, [
        "CPI→RMW は、本研究で最も具体的に書かれている重点仮説であり、専用の図表セットを先に用意する価値が高い。",
        "Gu, Kelly, Xiu (2020) の 0.40% は月次 stock-level OOS R²であり、『40%のリスク』とは異なる。",
        "長期データ取得と再走が終わると、因果性検定・IRF・FEVD・OOS評価の全体が一気につながりやすくなる。",
    ])

    doc.save(DOCX_OUTPUT_PATH)


def build_code_snapshot():
    text = dedent(
        f"""\
        from pathlib import Path

        MASTER_PLAN_PATH = Path(r"{MASTER_PLAN_PATH}")
        PPT_OUTPUT_PATH = Path(r"{PPT_OUTPUT_PATH}")
        XLSX_OUTPUT_PATH = Path(r"{XLSX_OUTPUT_PATH}")
        DOCX_OUTPUT_PATH = Path(r"{DOCX_OUTPUT_PATH}")

        THESIS_TITLE = {THESIS_TITLE!r}

        CURRENT_PROGRESS = {CURRENT_PROGRESS!r}
        CURRENT_ISSUES = {CURRENT_ISSUES!r}
        ROADMAP = {ROADMAP!r}

        def print_report():
            print(THESIS_TITLE)
            print()
            print("Current progress:")
            for item in CURRENT_PROGRESS:
                print("-", item)
            print()
            print("Current issues:")
            for item in CURRENT_ISSUES:
                print("-", item)
            print()
            print("Roadmap:")
            for idx, item in enumerate(ROADMAP, start=1):
                print(f"{{idx}}.", item)

        if __name__ == "__main__":
            print_report()
        """
    )
    Path(CODE_OUTPUT_PATH).write_text(text, encoding="utf-8")


def clear_attrs(path):
    subprocess.run(["xattr", "-c", str(path)], check=False)


def build_powerpoint():
    ppt_builder.OUT_PATH = Path(PPT_OUTPUT_PATH)
    ppt_builder.build_deck()


def build_all():
    Path(RESEARCH_DIR).mkdir(parents=True, exist_ok=True)
    build_powerpoint()
    build_excel()
    build_consolidated_docx()
    build_code_snapshot()
    clear_attrs(PPT_OUTPUT_PATH)
    clear_attrs(XLSX_OUTPUT_PATH)
    clear_attrs(DOCX_OUTPUT_PATH)
    clear_attrs(CODE_OUTPUT_PATH)
    print(PPT_OUTPUT_PATH)
    print(XLSX_OUTPUT_PATH)
    print(DOCX_OUTPUT_PATH)
    print(CODE_OUTPUT_PATH)


if __name__ == "__main__":
    build_all()
